from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

from ..utils.file_utils import ensure_directory_exists


class DocxGenerator:
    """Create professional DOCX deliverables for RegionGuard engagements."""

    def __init__(self, output_dir: Path):
        self.output_dir = ensure_directory_exists(output_dir)

    def generate_all(self, render_data: dict, base_name: str) -> list[Path]:
        files = [
            self.generate_authorization_letter(render_data, self.output_dir / f"{base_name}_authorization.docx"),
            self.generate_rules_of_engagement(render_data, self.output_dir / f"{base_name}_roe.docx"),
            self.generate_scope_definition(render_data, self.output_dir / f"{base_name}_scope.docx"),
        ]
        return files

    def generate_authorization_letter(self, render_data: dict, output_path: Path) -> Path:
        document = self._create_document()
        self._add_heading(document, "Engagement Authorization")
        self._add_key_value_table(
            document,
            [
                ("Client", render_data["client_name"]),
                ("Date prepared", date.today().strftime("%B %d, %Y")),
                ("Engagement type", render_data["engagement_preset_name"]),
                ("Target systems", render_data["target_type"]),
                ("Testing location", render_data["jurisdiction_name"]),
                ("Cloud provider", render_data["cloud_provider_name"]),
            ],
        )

        self._add_paragraph(
            document,
            "This document confirms that the client has approved the security assessment to proceed within the described scope and timeline.",
        )
        self._add_section(document, "Approval", [
            "The client authorizes testing of the in-scope systems listed below.",
            "Testing will occur during the approved timeframe and will respect the exclusions stated.",
        ])
        self._add_section(document, "Assessment goals", render_data.get("objectives_list", [render_data.get("objectives_text", "")]))
        self._add_section(document, "Systems in scope", render_data.get("scope_assets_list", [render_data.get("scope_text", "")]))

        if render_data.get("operational_notes"):
            self._add_section(document, "Special instructions", [render_data["operational_notes"]])

        self._add_paragraph(document, "Authorized testing", bold=True)
        self._add_section(document, "", [
            "Active security testing of approved systems and networks.",
            "Configuration and access control review.",
            "Reporting of findings with context and recommendations.",
        ])
        self._add_paragraph(document, "Signature", bold=True)
        self._add_paragraph(document, "Client authorized representative: _________________________________")
        self._add_paragraph(document, "Assessment lead: _________________________________")
        self._add_paragraph(document, "Date: _________________________________")

        document.save(output_path)
        return output_path

    def generate_rules_of_engagement(self, render_data: dict, output_path: Path) -> Path:
        document = self._create_document()
        self._add_heading(document, "Rules of Engagement")
        self._add_paragraph(document, "This document outlines how the assessment will be conducted and what to expect.")

        self._add_section(document, "What we will test", [
            "Active testing of authorized systems and network boundaries.",
            "Configuration review and access control verification.",
            "Documentation of findings with clear, actionable recommendations.",
        ])

        self._add_section(document, "What we will not do", [
            "Test outside the defined scope.",
            "Perform denial-of-service or sustained load testing (unless separately approved).",
            "Conduct social engineering or physical security testing.",
        ])

        if render_data.get("preset_roe_notes"):
            self._add_section(document, "Special guidance", render_data["preset_roe_notes"])

        self._add_section(document, "Communication during testing", [
            "The team will be available during the testing window for operational questions.",
            "Critical findings will be reported immediately.",
            "We will coordinate changes that might impact the business.",
        ])

        if render_data.get("testing_window"):
            self._add_section(document, "Testing window", [render_data["testing_window"]])

        if render_data.get("preset_testing_window"):
            self._add_section(document, "Recommended timing", [render_data["preset_testing_window"]])

        if render_data.get("preset_operational_considerations"):
            self._add_section(document, "Key things to know", render_data["preset_operational_considerations"])

        self._add_paragraph(document, "Questions?", bold=True)
        self._add_paragraph(document, "Ask the assessment lead before testing begins if anything is unclear.")
        document.save(output_path)
        return output_path

    def generate_scope_definition(self, render_data: dict, output_path: Path) -> Path:
        document = self._create_document()
        self._add_heading(document, "Scope Definition")
        self._add_paragraph(document, "This is what we're testing and what we're not testing.")

        self._add_key_value_table(
            document,
            [
                ("Client", render_data["client_name"]),
                ("Assessment type", render_data["engagement_preset_name"]),
                ("Target description", render_data["target_type"]),
                ("Testing location", render_data["jurisdiction_name"]),
                ("Cloud provider", render_data["cloud_provider_name"]),
            ],
        )

        if render_data.get("objectives_list"):
            self._add_section(document, "What we're assessing", render_data["objectives_list"])

        if render_data.get("scope_assets_list"):
            self._add_section(document, "What will be tested", render_data["scope_assets_list"])

        if render_data.get("exclusions_list"):
            self._add_section(document, "Out of scope (do not test)", render_data["exclusions_list"])

        scope_metadata = []
        scope_metadata.append(f"Testing window: {render_data['testing_window']}")
        if render_data['production_environment']:
            scope_metadata.append("This includes live production systems with real customer data.")
        else:
            scope_metadata.append("This is testing against staging or pre-production systems.")
        if render_data['authentication_provided']:
            scope_metadata.append("Credentials or API keys will be provided.")
        else:
            scope_metadata.append("Assessment is external/unauthenticated.")
        self._add_section(document, "Key details", scope_metadata)

        if render_data.get("operational_notes"):
            self._add_section(document, "Special instructions", [render_data["operational_notes"]])

        self._add_section(document, f"{render_data['cloud_provider_name']} context", [
            render_data.get("cloud_provider_summary", ""),
        ])

        provider_controls = render_data.get("cloud_provider_controls", [])
        if provider_controls:
            self._add_section(document, "Cloud built-in controls", provider_controls)

        self._add_paragraph(document, "This scope is designed for operational clarity. Review before client handoff.")
        document.save(output_path)
        return output_path

    def _create_document(self) -> Document:
        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        return document

    def _add_heading(self, document: Document, text: str, level: int = 1) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(18 if level == 1 else 14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.space_after = Pt(12)

    def _add_paragraph(self, document: Document, text: str, bold: bool = False) -> None:
        paragraph = document.add_paragraph(text)
        if paragraph.runs:
            paragraph.runs[0].bold = bold
        paragraph.space_after = Pt(8)

    def _add_section(self, document: Document, heading: str, points: list[str]) -> None:
        if heading:
            self._add_paragraph(document, heading, bold=True)
        for item in points:
            if item:
                self._add_bullet(document, item)

    def _add_bullet(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph(text, style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(4)

    def _add_key_value_table(self, document: Document, rows: list[tuple[str, str]]) -> None:
        table = document.add_table(rows=0, cols=2)
        table.style = "LightShading-Accent1"
        for label, value in rows:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value
            if row[0].paragraphs[0].runs:
                row[0].paragraphs[0].runs[0].bold = True
        table.autofit = True
        document.add_paragraph()
